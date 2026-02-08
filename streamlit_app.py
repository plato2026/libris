"""
LIBRIS - Advanced Librarian AI Agent
Streamlit Web Application for Public Deployment

FREE for the world to use!
Powered by Anthropic's Claude AI

VERSION 1.1 - Now with PDF and Word Document Support!
"""

import streamlit as st
import anthropic
import os
from datetime import datetime

# ============================================================================
# NEW: Import libraries for PDF and Word document processing
# ============================================================================
from pypdf import PdfReader
from docx import Document
import io

# ============================================================================
# PAGE CONFIGURATION
# ============================================================================

st.set_page_config(
    page_title="LIBRIS - Advanced Librarian AI",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ============================================================================
# LIBRIS SYSTEM PROMPT
# ============================================================================

LIBRIS_SYSTEM_PROMPT = """You are LIBRIS, an expert librarian and document analysis system specializing in historical and philosophical collections.

CORE CAPABILITIES:
1. Document Processing: Extract bibliographic data from uploaded documents
2. Intelligent Search: Search across ~1,100 historical/philosophical works
3. Thematic Analysis: Identify patterns and connections across texts
4. Export Formats: Provide results as BibTeX, CSV, JSON, or plain text

DOCUMENT PROCESSING PROTOCOL:
When a user uploads a document:
1. Extract bibliographic data (author, title, date, themes)
2. Structure into standardized entries
3. Integrate with existing knowledge
4. Provide processing statistics

Always respond with:
📚 **Document Processing Complete**
**File:** [filename]
**Entries Extracted:** [X] works
**Date Range:** [earliest] to [latest]
**Primary Themes:** [list themes]

SEARCH PROTOCOL:
Return results in markdown table format:

| Publication Date | Author | Book Title | Key Themes / Notes | Source |
|-----------------|--------|------------|-------------------|--------|
| [date] | [author] | [title] | [themes] | 📚/📄 |

Source indicators:
- 📚 Base = From LIBRIS core knowledge
- 📄 User Doc = From uploaded documents

ANALYSIS:
After each search, provide:
- Patterns observed (chronological, thematic)
- Insights from uploaded documents
- Suggested next steps
- Export options

TONE:
Professional but approachable, like a knowledgeable university librarian. Be precise, transparent about limitations, and enthusiastic about intellectual connections.

SPECIAL FEATURES:
- Transliteration-aware (match "Confucius" with "Kong Fuzi")
- Conceptual search (match "justice" with "dharma", "dikaiosyne")
- Multi-lingual titles (show original and translation)
- Cross-cultural perspectives

REMEMBER: You're helping make knowledge accessible to the world. Be helpful, educational, and inclusive of all intellectual traditions.
"""

# ============================================================================
# NEW: Document Processing Functions
# ============================================================================

def extract_text_from_pdf(uploaded_file):
    """
    Extract text from a PDF file.
    
    Args:
        uploaded_file: Streamlit UploadedFile object
        
    Returns:
        str: Extracted text from the PDF
    """
    try:
        # Read PDF
        pdf_reader = PdfReader(uploaded_file)
        
        # Extract text from all pages
        text = ""
        for page_num, page in enumerate(pdf_reader.pages, 1):
            page_text = page.extract_text()
            if page_text:
                text += f"\n--- Page {page_num} ---\n"
                text += page_text
        
        if not text.strip():
            return "⚠️ Could not extract text from PDF. The PDF might be image-based or encrypted."
        
        return text
    
    except Exception as e:
        return f"⚠️ Error reading PDF: {str(e)}"


def extract_text_from_docx(uploaded_file):
    """
    Extract text from a Word document (.docx).
    
    Args:
        uploaded_file: Streamlit UploadedFile object
        
    Returns:
        str: Extracted text from the document
    """
    try:
        # Read Word document
        doc = Document(uploaded_file)
        
        # Extract text from all paragraphs
        text = ""
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
        
        if not text.strip():
            return "⚠️ Could not extract text from Word document. The document might be empty."
        
        return text
    
    except Exception as e:
        return f"⚠️ Error reading Word document: {str(e)}"


def process_uploaded_file(uploaded_file):
    """
    Process an uploaded file based on its type.
    
    Args:
        uploaded_file: Streamlit UploadedFile object
        
    Returns:
        str: Extracted text content
    """
    filename = uploaded_file.name.lower()
    
    # PDF files
    if filename.endswith('.pdf'):
        return extract_text_from_pdf(uploaded_file)
    
    # Word documents
    elif filename.endswith('.docx'):
        return extract_text_from_docx(uploaded_file)
    
    # Text-based files (txt, md, csv)
    elif filename.endswith(('.txt', '.md', '.csv')):
        try:
            # Try UTF-8 first
            content = uploaded_file.read().decode('utf-8')
            return content
        except UnicodeDecodeError:
            # Fallback to latin-1
            uploaded_file.seek(0)  # Reset file pointer
            content = uploaded_file.read().decode('latin-1')
            return content
    
    else:
        return f"⚠️ Unsupported file type: {filename}"

# ============================================================================
# SESSION STATE INITIALIZATION
# ============================================================================

def init_session_state():
    """Initialize session state variables"""
    if 'messages' not in st.session_state:
        st.session_state.messages = []
    if 'documents' not in st.session_state:
        st.session_state.documents = []
    if 'api_key' not in st.session_state:
        st.session_state.api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if 'conversation_count' not in st.session_state:
        st.session_state.conversation_count = 0

# ============================================================================
# ANTHROPIC API FUNCTIONS
# ============================================================================

def chat_with_libris(user_message, api_key):
    """Send message to LIBRIS and get response"""
    try:
        client = anthropic.Anthropic(api_key=api_key)
        
        # Build message history
        messages = st.session_state.messages + [
            {"role": "user", "content": user_message}
        ]
        
        # Call Claude API
        response = client.messages.create(
            model="claude-sonnet-4-5-20250929",
            max_tokens=4000,
            system=LIBRIS_SYSTEM_PROMPT,
            messages=messages
        )
        
        assistant_message = response.content[0].text
        
        # Update conversation history
        st.session_state.messages.append({"role": "user", "content": user_message})
        st.session_state.messages.append({"role": "assistant", "content": assistant_message})
        st.session_state.conversation_count += 1
        
        return assistant_message
        
    except anthropic.AuthenticationError:
        return "❌ **Authentication Error**: Invalid API key. Please check your API key in the sidebar."
    except anthropic.RateLimitError:
        return "⚠️ **Rate Limit**: Too many requests. Please wait a moment and try again."
    except Exception as e:
        return f"❌ **Error**: {str(e)}"

# ============================================================================
# UI COMPONENTS
# ============================================================================

def render_header():
    """Render the app header"""
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        st.markdown("""
        <h1 style='text-align: center; color: #6366f1;'>
            📚 LIBRIS
        </h1>
        <p style='text-align: center; font-size: 1.2em; color: #94a3b8;'>
            Advanced Librarian AI Agent
        </p>
        <p style='text-align: center; color: #64748b;'>
            Free for the world to use • Powered by Claude AI
        </p>
        """, unsafe_allow_html=True)

def render_sidebar():
    """Render the sidebar with info and stats"""
    with st.sidebar:
        st.markdown("### 🎯 About LIBRIS")
        st.markdown("""
        LIBRIS specializes in:
        - 📄 Document processing
        - 🔍 Intelligent search
        - 🌍 Cross-cultural perspectives
        - 📚 ~1,100 historical & philosophical works
        """)
        
        st.markdown("---")
        
        st.markdown("### 📊 Your Session")
        st.metric("Documents Processed", len(st.session_state.documents))
        st.metric("Queries Made", st.session_state.conversation_count)
        
        st.markdown("---")
        
        st.markdown("### 🔑 API Configuration")
        
        # Check if API key exists
        if st.session_state.api_key:
            st.success("✅ API key configured")
            if st.button("🔄 Reset Session"):
                st.session_state.messages = []
                st.session_state.documents = []
                st.session_state.conversation_count = 0
                st.rerun()
        else:
            st.warning("⚠️ No API key configured")
            st.markdown("""
            **For Administrators:**
            Set `ANTHROPIC_API_KEY` in Streamlit secrets.
            
            **For Local Testing:**
            Set environment variable or enter key below (not recommended for production).
            """)
            
            temp_key = st.text_input("Temporary API Key (testing only)", type="password")
            if temp_key:
                st.session_state.api_key = temp_key
                st.rerun()
        
        st.markdown("---")
        
        st.markdown("### 💡 Example Queries")
        examples = [
            "Ancient Greek philosophy",
            "Social contract theory",
            "Buddhist ethics",
            "Medieval Islamic philosophy",
            "Confucian virtue ethics",
            "Natural law tradition"
        ]
        
        for example in examples:
            st.markdown(f"- {example}")
        
        st.markdown("---")
        
        st.markdown("### ℹ️ How to Use")
        st.markdown("""
        1. **Search**: Enter queries in the search box
        2. **Upload**: Upload reading lists or bibliographies
        3. **Chat**: Ask questions about texts
        4. **Export**: Request BibTeX, CSV, or JSON
        """)
        
        st.markdown("---")
        
        st.markdown("""
        <div style='text-align: center; font-size: 0.8em; color: #64748b;'>
        Made with ❤️ for the world<br>
        Open source • Free forever<br>
        v1.1 - PDF & Word Support
        </div>
        """, unsafe_allow_html=True)

def render_welcome():
    """Render welcome message when no conversation exists"""
    st.markdown("""
    <div style='padding: 2rem; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
                border-radius: 10px; color: white; text-align: center; margin: 2rem 0;'>
        <h2>👋 Welcome to LIBRIS!</h2>
        <p style='font-size: 1.1em;'>Your advanced AI librarian for historical and philosophical research</p>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("""
        ### 🔍 **Search**
        Find books and texts across thousands of years of human thought. Search by:
        - Author names
        - Time periods
        - Themes & concepts
        - Cultural traditions
        """)
    
    with col2:
        st.markdown("""
        ### 📄 **Process Documents**
        Upload your reading lists, syllabi, or bibliographies. LIBRIS will:
        - Extract bibliographic data
        - Categorize by theme
        - Identify patterns
        - Build your collection
        """)
    
    with col3:
        st.markdown("""
        ### 📚 **Learn & Explore**
        Discover connections across:
        - Ancient Greek & Roman
        - Islamic Golden Age
        - Chinese classics
        - Indian philosophy
        - Modern thought
        """)

# ============================================================================
# MAIN APP
# ============================================================================

def main():
    """Main application logic"""
    
    # Initialize session state
    init_session_state()
    
    # Render header and sidebar
    render_header()
    render_sidebar()
    
    # Check for API key
    if not st.session_state.api_key:
        st.error("⚠️ **API Key Required**: LIBRIS requires an Anthropic API key to function. Please configure it in the sidebar.")
        st.info("💡 **Note for Users**: If you're seeing this on a public deployment, the administrator needs to configure the API key in Streamlit Cloud secrets.")
        return
    
    # Create tabs
    tab1, tab2, tab3, tab4 = st.tabs(["🔍 Search", "📄 Upload Document", "💬 Chat", "📊 Export"])
    
    # TAB 1: SEARCH
    with tab1:
        st.markdown("### 🔍 Search LIBRIS Knowledge Base")
        
        col1, col2 = st.columns([4, 1])
        with col1:
            search_query = st.text_input(
                "Search Query",
                placeholder="e.g., 'ancient Greek ethics', 'works by Plato', '18th century social contract'",
                label_visibility="collapsed"
            )
        with col2:
            search_button = st.button("🔍 Search", use_container_width=True)
        
        if search_button and search_query:
            with st.spinner("🔍 Searching LIBRIS knowledge base..."):
                response = chat_with_libris(f"Search for: {search_query}", st.session_state.api_key)
                st.markdown(response)
        
        # Quick search buttons
        st.markdown("**Quick searches:**")
        quick_cols = st.columns(3)
        quick_searches = [
            "Ancient Greek philosophy",
            "Social contract theory", 
            "Buddhist ethics",
            "Medieval Islamic philosophy",
            "Confucian virtue",
            "Natural law"
        ]
        
        for idx, qs in enumerate(quick_searches):
            with quick_cols[idx % 3]:
                if st.button(qs, key=f"quick_{idx}"):
                    with st.spinner(f"Searching for {qs}..."):
                        response = chat_with_libris(f"Search for: {qs}", st.session_state.api_key)
                        st.markdown(response)
    
    # TAB 2: UPLOAD DOCUMENT (UPDATED!)
    with tab2:
        st.markdown("### 📄 Upload Document for Processing")
        
        # UPDATED: Now accepts PDF and Word documents!
        uploaded_file = st.file_uploader(
            "Choose a file",
            type=['txt', 'md', 'csv', 'pdf', 'docx'],  # UPDATED: Added 'pdf' and 'docx'
            help="Upload reading lists, syllabi, bibliographies, PDFs, or Word documents"
        )
        
        if uploaded_file is not None:
            # Process the file based on type
            with st.spinner(f"Reading {uploaded_file.name}..."):
                content = process_uploaded_file(uploaded_file)
            
            # Check if extraction was successful
            if content.startswith("⚠️"):
                st.warning(content)
                st.info("💡 **Tip**: For image-based PDFs, try converting to text first using an OCR tool.")
            else:
                st.success(f"✅ File loaded: {uploaded_file.name}")
                
                # Show file type info
                if uploaded_file.name.lower().endswith('.pdf'):
                    st.info("📄 **PDF Document** - Text extracted from all pages")
                elif uploaded_file.name.lower().endswith('.docx'):
                    st.info("📝 **Word Document** - Text extracted from paragraphs and tables")
                
                with st.expander("Preview file content"):
                    preview_length = 2000
                    if len(content) > preview_length:
                        st.text(content[:preview_length] + f"\n\n... ({len(content) - preview_length} more characters)")
                    else:
                        st.text(content)
                
                if st.button("📚 Process Document", type="primary"):
                    with st.spinner(f"Processing {uploaded_file.name}..."):
                        message = f"I'm uploading a document called '{uploaded_file.name}'. Please process it and extract bibliographic information.\n\nDocument content:\n{content}"
                        response = chat_with_libris(message, st.session_state.api_key)
                        
                        # Store document info
                        st.session_state.documents.append({
                            'filename': uploaded_file.name,
                            'processed_at': datetime.now().isoformat(),
                            'file_type': uploaded_file.type
                        })
                        
                        st.markdown(response)
        
        st.markdown("---")
        st.markdown("""
        **Supported formats:**
        - 📄 `.pdf` - PDF documents (text-based)
        - 📝 `.docx` - Word documents
        - 📝 `.txt` - Plain text files
        - 📝 `.md` - Markdown files
        - 📊 `.csv` - CSV files
        
        **What to upload:**
        - Course syllabi (PDF or Word)
        - Reading lists
        - Bibliographies
        - Research papers
        - Research notes
        
        **Note:** For best results with PDFs, ensure they are text-based (not scanned images).
        """)
    
    # TAB 3: CHAT
    with tab3:
        st.markdown("### 💬 Chat with LIBRIS")
        
        # Display chat history
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
        
        # Chat input
        if prompt := st.chat_input("Ask LIBRIS anything about historical or philosophical texts..."):
            # Display user message
            with st.chat_message("user"):
                st.markdown(prompt)
            
            # Get and display assistant response
            with st.chat_message("assistant"):
                with st.spinner("LIBRIS is thinking..."):
                    response = chat_with_libris(prompt, st.session_state.api_key)
                    st.markdown(response)
    
    # TAB 4: EXPORT
    with tab4:
        st.markdown("### 📊 Export Search Results")
        
        st.markdown("""
        Export your last search results in various academic formats.
        First perform a search, then come here to export the results.
        """)
        
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("📑 Export as BibTeX", use_container_width=True):
                with st.spinner("Generating BibTeX..."):
                    response = chat_with_libris("Export the last results as BibTeX", st.session_state.api_key)
                    st.markdown(response)
            
            if st.button("📄 Export as JSON", use_container_width=True):
                with st.spinner("Generating JSON..."):
                    response = chat_with_libris("Export the last results as JSON", st.session_state.api_key)
                    st.markdown(response)
        
        with col2:
            if st.button("📊 Export as CSV", use_container_width=True):
                with st.spinner("Generating CSV..."):
                    response = chat_with_libris("Export the last results as CSV", st.session_state.api_key)
                    st.markdown(response)
            
            if st.button("📝 Export as Plain Text", use_container_width=True):
                with st.spinner("Generating plain text..."):
                    response = chat_with_libris("Export the last results as plain text", st.session_state.api_key)
                    st.markdown(response)
    
    # Show welcome message if no conversation
    if len(st.session_state.messages) == 0:
        st.markdown("---")
        render_welcome()

# ============================================================================
# RUN APP
# ============================================================================

if __name__ == "__main__":
    main()
